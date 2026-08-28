"""
Generación automática del Anexo II: esquema gráfico del edificio con la
posición de cada detector marcada sobre el plano correspondiente.

Los datos se leen directamente del Excel original:
  - Hoja "Planos": columna A = nombre del plano (o "(Foto exterior del
    centro)"), columna B = imagen embebida (una foto/plano por fila).
  - Hoja "Detectores": columnas "Nombre del plano", "Punto X", "Punto Y"
    (coordenadas relativas 0-1 sobre la imagen del plano), "Código"
    (código del detector) y "Código de la sala".

Por cada plano distinto que tenga al menos un detector asociado se genera
una página con:
  - el plano, con un punto rojo en cada posición exacta (las coordenadas
    relativas no se alteran; solo se recolocan las ETIQUETAS de texto para
    que no se solapen entre sí),
  - el código del detector junto al punto y, debajo, entre paréntesis, el
    código de la sala,
  - el logotipo y la foto exterior del centro en una columna aparte, sin
    tapar el plano.
"""

from __future__ import annotations

import io
import unicodedata

import openpyxl
import pandas as pd
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt
from PIL import Image, ImageDraw, ImageFont

from utils_informe.assets import logo_por_defecto

PUNTO_RADIO_FRAC = 0.010  # radio del punto, como fracción del lado menor de la imagen
MARGEN_FRAC = 0.012


def _normalizar(texto) -> str:
    texto = str(texto or "").strip().lower()
    return "".join(c for c in unicodedata.normalize("NFD", texto) if unicodedata.category(c) != "Mn")


class Anexo2Error(ValueError):
    """No se han podido extraer los datos necesarios del Excel para generar el Anexo II."""


def extraer_datos_planos(file) -> dict:
    """Lee el Excel original (ruta o fichero) y devuelve:
    {
        "planos": {nombre_normalizado: {"nombre": str, "imagen": bytes}},
        "foto_exterior": bytes | None,
        "puntos": {nombre_normalizado: [{"x": float, "y": float, "codigo": str, "sala": str}]},
    }
    """
    if hasattr(file, "seek"):
        file.seek(0)
    wb = openpyxl.load_workbook(file, data_only=True)

    if "Planos" not in wb.sheetnames or "Detectores" not in wb.sheetnames:
        raise Anexo2Error("El Excel no tiene las hojas 'Planos' y 'Detectores' necesarias.")

    ws_planos = wb["Planos"]

    # nombre (columna A) de cada fila que tiene una imagen anclada en columna B
    nombres_por_fila = {}
    for row in ws_planos.iter_rows(min_col=1, max_col=1):
        for cell in row:
            if cell.value is not None:
                nombres_por_fila[cell.row] = str(cell.value).strip()

    planos = {}
    foto_exterior = None
    for img in ws_planos._images:
        fila_excel = img.anchor._from.row + 1  # openpyxl es 0-index en el anchor
        nombre = nombres_por_fila.get(fila_excel, f"Plano fila {fila_excel}")
        data = img._data()
        clave = _normalizar(nombre)
        if "exterior" in clave:
            foto_exterior = data
        else:
            planos[clave] = {"nombre": nombre, "imagen": data}

    if not planos:
        raise Anexo2Error("No se ha encontrado ninguna imagen de plano en la hoja 'Planos'.")

    # puntos por plano, desde la hoja Detectores (fila 2 = cabecera real)
    ws_det = wb["Detectores"]
    raw = list(ws_det.iter_rows(values_only=True))
    header = raw[1]
    col_idx = {str(h).strip(): i for i, h in enumerate(header) if h is not None}

    requeridas = ["Nombre del plano", "Punto X", "Punto Y", "Código", "Código de la sala"]
    faltan = [c for c in requeridas if c not in col_idx]
    if faltan:
        raise Anexo2Error(
            "Faltan columnas en la hoja 'Detectores' para generar el Anexo II: " + ", ".join(faltan)
        )

    puntos: dict[str, list[dict]] = {}
    for fila in raw[2:]:
        if fila is None:
            continue
        nombre_plano = fila[col_idx["Nombre del plano"]]
        x = fila[col_idx["Punto X"]]
        y = fila[col_idx["Punto Y"]]
        if nombre_plano is None or x is None or y is None:
            continue
        codigo = fila[col_idx["Código"]] or ""
        sala_cod = fila[col_idx["Código de la sala"]] or ""
        clave = _normalizar(nombre_plano)
        puntos.setdefault(clave, []).append(
            {"x": float(x), "y": float(y), "codigo": str(codigo).strip(), "sala": str(sala_cod).strip()}
        )

    return {"planos": planos, "foto_exterior": foto_exterior, "puntos": puntos}


def _cargar_fuente(tamano: int):
    try:
        return ImageFont.load_default(size=tamano)
    except TypeError:
        # Pillow < 10.1: load_default() no admite tamaño variable
        return ImageFont.load_default()


def _bbox_multilinea(draw, xy, lineas, font, spacing=2):
    return draw.multiline_textbbox(xy, "\n".join(lineas), font=font, spacing=spacing, align="center")


def _rectangulos_solapan(a, b, margen=2):
    return not (
        a[2] + margen < b[0]
        or a[0] - margen > b[2]
        or a[3] + margen < b[1]
        or a[1] - margen > b[3]
    )


def componer_plano(imagen_bytes: bytes, puntos: list[dict]) -> bytes:
    """Dibuja los puntos rojos y las etiquetas (código de detector y, debajo
    entre paréntesis, código de sala) sobre la imagen del plano, evitando
    que las etiquetas se solapen entre sí. La posición del punto en sí
    nunca se modifica: solo se recoloca el texto."""
    img = Image.open(io.BytesIO(imagen_bytes)).convert("RGB")
    w, h = img.size
    draw = ImageDraw.Draw(img)

    lado_menor = min(w, h)
    radio = max(4, int(lado_menor * PUNTO_RADIO_FRAC))
    margen = max(4, int(lado_menor * MARGEN_FRAC))
    font = _cargar_fuente(max(12, int(lado_menor * 0.032)))

    cajas_ocupadas: list[tuple[float, float, float, float]] = []

    # candidatos de desplazamiento (dx, dy) respecto al punto, de más cercano
    # a más lejano; se prueban en este orden hasta encontrar uno libre
    def candidatos(bw, bh):
        pad = radio + 4
        base = [
            (pad, -bh / 2),  # derecha
            (-bw - pad, -bh / 2),  # izquierda
            (-bw / 2, -bh - pad),  # arriba
            (-bw / 2, pad),  # abajo
            (pad, -bh - pad),  # arriba-derecha
            (pad, pad),  # abajo-derecha
            (-bw - pad, -bh - pad),  # arriba-izquierda
            (-bw - pad, pad),  # abajo-izquierda
        ]
        for factor in (1.0, 1.8, 2.6, 3.4, 4.2):
            for dx, dy in base:
                yield dx * factor if dx > 0 else (dx - pad) * (factor - 1) + dx, dy * factor if dy > 0 else (dy - pad) * (factor - 1) + dy

    for punto in puntos:
        px = punto["x"] * w
        py = punto["y"] * h

        # punto rojo (posición fija, no se mueve nunca)
        draw.ellipse([px - radio, py - radio, px + radio, py + radio], fill=(220, 0, 0), outline=(90, 0, 0))

        lineas = [punto["codigo"]] if punto["codigo"] else []
        if punto["sala"]:
            lineas.append(f"({punto['sala']})")
        if not lineas:
            continue

        bbox_0 = _bbox_multilinea(draw, (0, 0), lineas, font)
        bw = bbox_0[2] - bbox_0[0]
        bh = bbox_0[3] - bbox_0[1]

        colocado = False
        for dx, dy in candidatos(bw, bh):
            cx, cy = px + dx, py + dy
            caja = (cx, cy, cx + bw, cy + bh)
            dentro_imagen = caja[0] >= 0 and caja[1] >= 0 and caja[2] <= w and caja[3] <= h
            if not dentro_imagen:
                continue
            if any(_rectangulos_solapan(caja, otra) for otra in cajas_ocupadas):
                continue
            colocado = True
            break
        if not colocado:
            # no se encontró hueco libre: se deja a la derecha del punto,
            # recortando si hiciera falta, para no perder la etiqueta
            cx, cy = min(px + radio + 4, w - bw - 1), max(0, py - bh / 2)

        # Texto en rojo, sin ningún recuadro/fondo detrás (a petición
        # expresa: antes llevaba una caja blanca con borde rojo).
        draw.multiline_text((cx, cy), "\n".join(lineas), font=font, fill=(200, 0, 0), spacing=2, align="center")
        cajas_ocupadas.append((cx - 2, cy - 1, cx + bw + 2, cy + bh + 1))

    buffer = io.BytesIO()
    img.save(buffer, format="PNG")
    return buffer.getvalue()


def _fit(img_bytes: bytes, max_width_cm: float, max_height_cm: float):
    img = Image.open(io.BytesIO(img_bytes))
    w_px, h_px = img.size
    aspect = w_px / h_px if h_px else 1
    max_w, max_h = Cm(max_width_cm), Cm(max_height_cm)
    height = max_h
    width = int(height * aspect)
    if width > max_w:
        width = max_w
        height = int(width / aspect)
    return width, height


def generar_documento_anexo2(
    datos: dict,
    logo_bytes: bytes | None = None,
    centro: str = "",
    titulo: str = "ANEXO II: ESQUEMA GRÁFICO DO EDIFICIO E PLANOS DE CADA PLANTA",
) -> bytes:
    """Genera un .docx con una página (en horizontal) por cada plano que
    tenga detectores asociados: el plano con los puntos marcados a la
    izquierda, y el logotipo + foto exterior del centro (con su nombre
    debajo) a la derecha, sin tapar el plano."""
    planos = datos["planos"]
    puntos_por_plano = datos["puntos"]
    foto_exterior = datos.get("foto_exterior")
    logo_a_usar = logo_bytes or logo_por_defecto()

    nombres_con_puntos = [clave for clave in planos if clave in puntos_por_plano and puntos_por_plano[clave]]
    if not nombres_con_puntos:
        raise Anexo2Error(
            "Ningún plano tiene detectores con coordenadas (Punto X / Punto Y) asociadas."
        )

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(10.5)

    section = doc.sections[0]
    # Página en horizontal, tamaño A4 explícito (el tamaño por defecto de
    # python-docx es Letter, no A4).
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Cm(29.7)
    section.page_height = Cm(21.0)
    section.left_margin = Cm(1.3)
    section.right_margin = Cm(1.3)
    section.top_margin = Cm(1.3)
    section.bottom_margin = Cm(1.3)

    ancho_util_cm = (section.page_width - section.left_margin - section.right_margin) / 360000
    alto_util_cm = (section.page_height - section.top_margin - section.bottom_margin) / 360000
    col_lateral_cm = 6.0
    col_plano_cm = ancho_util_cm - col_lateral_cm - 0.3

    for i, clave in enumerate(nombres_con_puntos):
        if i > 0:
            doc.add_page_break()

        plano_info = planos[clave]
        p = doc.add_paragraph()
        run = p.add_run(f"{titulo} — {plano_info['nombre']}")
        run.bold = True
        run.font.size = Pt(12)
        p.paragraph_format.space_after = Pt(8)

        compuesto = componer_plano(plano_info["imagen"], puntos_por_plano[clave])

        table = doc.add_table(rows=1, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for col, w in zip(table.columns, (Cm(col_plano_cm), Cm(col_lateral_cm))):
            col.width = w

        celda_plano = table.rows[0].cells[0]
        celda_plano.width = Cm(col_plano_cm)
        p_plano = celda_plano.paragraphs[0]
        p_plano.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_plano = p_plano.add_run()
        w_plano, h_plano = _fit(compuesto, col_plano_cm - 0.3, alto_util_cm - 1.5)
        run_plano.add_picture(io.BytesIO(compuesto), width=w_plano, height=h_plano)

        celda_lateral = table.rows[0].cells[1]
        celda_lateral.width = Cm(col_lateral_cm)
        p_lat1 = celda_lateral.paragraphs[0]
        p_lat1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if logo_a_usar:
            r = p_lat1.add_run()
            w_logo, h_logo = _fit(logo_a_usar, col_lateral_cm - 0.6, 5.2)
            r.add_picture(io.BytesIO(logo_a_usar), width=w_logo, height=h_logo)
        else:
            p_lat1.add_run("[logotipo]").italic = True

        if foto_exterior:
            p_lat2 = celda_lateral.add_paragraph()
            p_lat2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_lat2.paragraph_format.space_before = Pt(14)
            r2 = p_lat2.add_run()
            w_foto, h_foto = _fit(foto_exterior, col_lateral_cm - 0.6, 6.5)
            r2.add_picture(io.BytesIO(foto_exterior), width=w_foto, height=h_foto)

            # nombre del centro, debajo de la foto exterior
            if centro:
                nombre_p = celda_lateral.add_paragraph()
                nombre_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                nombre_p.paragraph_format.space_before = Pt(2)
                nombre_run = nombre_p.add_run(centro)
                nombre_run.bold = True
                nombre_run.font.size = Pt(9.5)
        elif centro:
            # sin foto exterior disponible, se muestra igualmente el nombre
            nombre_p = celda_lateral.add_paragraph()
            nombre_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            nombre_p.paragraph_format.space_before = Pt(14)
            nombre_run = nombre_p.add_run(centro)
            nombre_run.bold = True
            nombre_run.font.size = Pt(9.5)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
