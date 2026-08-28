"""
Generacion del "Informe de resultados de medicions de radon no centro de
traballo" siguiendo la plantilla del Servizo Galego de Saude / UPRL, con la
documentacion de referencia do Consello de Seguridade Nuclear (CSN).

Genera un .docx en memoria (BytesIO) a partir de:
  - los datos de cabecera introducidos en el formulario (ReportContext)
  - la tabla de resultados de mediciones (DataFrame `df`, ya filtrado a un
    unico centro de traballo)
  - opcionalmente, imagenes de logotipo para la cabecera institucional

Soporta dos idiomas de salida ("gl" gallego, "es" castellano) mediante el
parametro `idioma` de `generate_report`. El contenido dinamico (nombres de
sala, categorias profesionales) debe venir ya en el idioma deseado en el
`ReportContext` (ver `utils.excel_parser`, parametro `traducir_galego`).
"""

from __future__ import annotations

import io
import re
from dataclasses import dataclass, field

import pandas as pd
from docx import Document
from docx.enum.table import WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from PIL import Image

from utils_informe.excel_parser import _strip_accents
from utils_informe.assets import logo_por_defecto

NIVEL_REFERENCIA_BQ_M3 = 300

AREAS_GRUPO_XERAL = {"", "atencion primaria", "pac", "atencion primaria + pac", "consultorio"}


def _area_grupo(area: str) -> str:
    clave = _strip_accents(area or "").strip().lower()
    return "A" if clave in AREAS_GRUPO_XERAL else "B"


def _tipo_zona_texto(tipo_zona: str, idioma: str) -> str:
    """Palabra a usar en vez de "unidade/servizo" (o "unidad/servicio")
    quando a área non é ningunha das excepcións xerais: a que se
    escolleu explicitamente ("unidade"/"servizo"), ou un oco en
    branco (".....") para completar a man se se escolleu "outro tipo
    de zona" ou non se escolleu nada."""
    if tipo_zona == "unidade":
        return "unidade" if idioma == "gl" else "unidad"
    if tipo_zona == "servizo":
        return "servizo" if idioma == "gl" else "servicio"
    return "....."


def _articulo_zona(tipo_zona: str, idioma: str) -> str:
    """Artigo (concordando en xénero) que vai antes de "unidade"
    (femenino: "da"/"de la") ou "servizo" (masculino: "do"/"del")."""
    if tipo_zona == "servizo":
        return "do" if idioma == "gl" else "del"
    return "da" if idioma == "gl" else "de la"


TEXTOS = {
    "gl": {
        "logo_placeholder": "UPRL · SERVIZO GALEGO DE SAÚDE · ÁREA SANITARIA",
        "data_label": "Data",
        "paxina_label": "Páxina",
        "de_label": "de",
        "titulo_grupo_a": "INFORME DE RESULTADOS DAS MEDICIÓNS DE RADON NO {centro}",
        "titulo_grupo_b": (
            "INFORME DE RESULTADOS DAS MEDICIÓNS DE RADON NA {tipo_zona_upper} DE "
            "{area}, NO {centro}"
        ),
        "heading_1": "IDENTIFICACIÓN DO CENTRO DE TRABALLO",
        "lbl_xerencia": "XERENCIA",
        "lbl_cif": "CIF",
        "lbl_centro": "CENTRO",
        "lbl_servizo": "SERVIZO / UNIDADE MOSTREXADA",
        "lbl_enderezo": "ENDEREZO",
        "lbl_descricion": "DESCRICIÓN DA ÁREA MOSTREXADA",
        "lbl_superficie_construida": "Superficie construída",
        "lbl_superficie_util": "Superficie útil",
        "lbl_num_plantas": "N.º plantas",
        "heading_2": "OBXECTO DO INFORME",
        "obxecto_lugar_b": "{tipo_zona_cap} de {area}, no {centro}",
        "obxecto_texto": (
            "O presente informe ten como obxecto recoller os resultados das medicións de "
            "radon no {lugar}, para, no caso de que a exposición a radon sexa superior ao "
            "nivel de referencia establecido ({nivel} Bq/m³), establecer as medidas de "
            "remediación necesarias, para diminuír ou eliminar estes niveis."
        ),
        "heading_3": "INFORMACIÓN SOBRE OS/AS TRABALLADORES/AS",
        "intro_suxeito_b": "{tipo_zona_cap} de {area} no {centro}",
        "intro_texto_a": (
            "En base á documentación e lexislación de referencia determináronse no {centro} "
            "as seguintes zonas de traballo susceptibles de exposición a radón pola súa "
            "situación en planta baixa ou baixo rasante, coa ocupación que se sinala a "
            "continuación:"
        ),
        "intro_texto_a_sen_ocupacion": (
            "En base á documentación e lexislación de referencia determináronse no {centro} "
            "as seguintes zonas de traballo susceptibles de exposición a radón pola súa "
            "situación en planta baixa ou baixo rasante:"
        ),
        "intro_texto": (
            "O {suxeito}, consta dos seguintes postos de traballo e a seguinte ocupación por "
            "espazo de traballo, susceptibles de exposición a radón pola súa situación en "
            "planta baixa ou baixo rasante:"
        ),
        "sin_dato": ".....",
        "num_traballadores_texto": (
            "O número de traballadores adscritos a este centro é de {total}, divididos nas "
            "seguintes categorías:"
        ),
        "cierre_grupo_a": "Os traballadores/as do {centro} foron informados da ",
        "cierre_grupo_b": (
            "Os traballadores/as {articulo} {tipo_zona} de {area} no {centro}, foron informados da "
        ),
        "cierre_resto": (
            "realización das medicións e da súa finalidade mediante {medio} o día {data}."
        ),
        "medio_defecto": "correo electrónico",
        "heading_4": "CONDICIÓNS DA EXPOSICIÓN",
        "texto_4": (
            "No formulario de toma de datos do Anexo I recóllense: o tempo de exposición "
            "de cada detector, o responsable da súa colocación e retirada, a comprobación do "
            "estado dos detectores, as condicións de temperatura ou humidade, e observacións "
            "sobre o sistema de ventilación."
        ),
        "texto_4_detectores": (
            "Descrición dos detectores utilizados: DETECTORES DE TRAZAS NUCLEARES EN ESTADO "
            "SOLIDO (SSNTD) DE LONGA DURACION. TIPO/MODELO: PADC/CR-39 – RSKS. FABRICANTE: "
            "RADOSYS-HUNGRIA."
        ),
        "heading_5": "PLANOS",
        "texto_5": (
            "O esquema gráfico do edificio e planos de cada planta móstranse no Anexo II, "
            "onde se indican as zonas de mostraxe e as localizacións dos detectores; xunto co "
            "código de cada zona de mostraxe, márcase o código do detector correspondente."
        ),
        "heading_7": "RESULTADOS DAS MEDICIÓNS REALIZADAS",
        "texto_7_intro": "Os resultados das medicións de radon realizadas poden verse na seguinte táboa:",
        "texto_7_sin_datos": "Non hai datos de mediciones cargados para este centro de traballo.",
        "tabla_headers": [
            "Código zona de mostraxe",
            "Código detector",
            "Data de inicio exposición",
            "Data fin de exposición",
            "Concentración radón (Bq/m³) (*)",
            "Incerteza expandida e K",
            "Sala / Posto de traballo asociado",
        ],
        "header_sala_sin_puestos": "Sala do posto de traballo asociado",
        "tabla_nota": "* Resáltanse en negriña e letra vermella os valores superiores a {nivel} Bq/m³.",
        "heading_8": "CONCLUSIÓNS",
        "conclusion_no_supera": (
            "A vista dos datos do apartado 7, non hai ningún posto de traballo nos que se "
            "supere o nivel de referencia."
        ),
        "conclusion_supera": (
            "As medicións de radon nos seguintes lugares de traballo: {salas}, dan valores "
            "superiores ao nivel de referencia. Para estes casos estudiaranse medidas de "
            "remediación, completaranse as medicións e/ou "
            "efectuaranse medicións en continuo, segundo proceda."
        ),
        "salas_indicadas": "indicados no apartado 7",
        "heading_9": "DOCUMENTACIÓN DE REFERENCIA",
        "referencias": [
            "Lei 31/1995, de 8 de novembro, de Prevención de Riscos Laborais (BOE nº 269, de 10 de novembro).",
            "Real Decreto 39/1997, de 17 de xaneiro, polo que se aproba o Regulamento dos servizos de "
            "prevención (BOE nº 27, de 31 de xaneiro).",
            "Real Decreto 732/2019, de 20 de decembro, polo que se modifica o Código Técnico da "
            "Edificación, aprobado polo Real Decreto 314/2006, de 17 de marzo (BOE nº 311, de 27 de "
            "decembro, páxinas 140488 a 140674).",
            "Instrución IS-47, de 9 de abril de 2025, do Consello de Seguridade Nuclear pola que se "
            "aproba o listado de términos municipais de actuación prioritaria contra o radón e establéce.",
            "Guía de Seguridade do CSN 11.4 Metodoloxía para a avaliación da exposición ao radon nos "
            "lugares de traballo.",
            "Estratexia para reducir a exposición ao radon en Galicia. Estratexia Reduce Radon 2025-2030. "
            "Consellería de Sanidade, Dirección Xeral de Saúde Pública, ano 2024. "
            "https://airesaude.sergas.gal/Contidos/Documents/135/Estratexia%20reduce%20radon%202025-2030.pdf",
        ],
        "heading_10": "DATA E FIRMA DO TÉCNICO/A",
        "data_label_10": "Data",
        "firma_label": "Asdo.",
        "anexos": [
            "ANEXO I: FORMULARIOS TOMA DE DATOS",
            "ANEXO II: ESQUEMA GRÁFICO DO EDIFICIO E PLANOS DE CADA PLANTA",
            "ANEXO III: INFORME DE ENSAIO DO LABORATORIO ACREDITADO",
            "ANEXO IV: CERTIFICADO ENAC DO LABORATORIO ACREDITADO",
        ],
        "placeholder_centro": "...................",
        "placeholder_area": "...................",
    },
    "es": {
        "logo_placeholder": "UPRL · SERVICIO GALLEGO DE SALUD · ÁREA SANITARIA",
        "data_label": "Fecha",
        "paxina_label": "Página",
        "de_label": "de",
        "titulo_grupo_a": "INFORME DE RESULTADOS DE MEDICIONES DE RADÓN EN EL {centro}",
        "titulo_grupo_b": (
            "INFORME DE RESULTADOS DE MEDICIONES DE RADÓN EN LA {tipo_zona_upper} DE "
            "{area}, EN EL {centro}"
        ),
        "heading_1": "IDENTIFICACIÓN DEL CENTRO DE TRABAJO",
        "lbl_xerencia": "GERENCIA",
        "lbl_cif": "CIF",
        "lbl_centro": "CENTRO",
        "lbl_servizo": "SERVICIO / UNIDAD MUESTREADA",
        "lbl_enderezo": "DIRECCIÓN",
        "lbl_descricion": "DESCRIPCIÓN DEL ÁREA MUESTREADA",
        "lbl_superficie_construida": "Superficie construida",
        "lbl_superficie_util": "Superficie útil",
        "lbl_num_plantas": "N.º plantas",
        "heading_2": "OBJETO DEL INFORME",
        "obxecto_lugar_b": "{tipo_zona_cap} de {area}, en el {centro}",
        "obxecto_texto": (
            "El presente informe tiene como objeto recoger los resultados de las mediciones "
            "de radón en el {lugar}, para, en caso de que la exposición a radón sea superior "
            "al nivel de referencia establecido ({nivel} Bq/m³), establecer las medidas de "
            "remediación necesarias, para disminuir o eliminar estos niveles."
        ),
        "heading_3": "INFORMACIÓN SOBRE LOS/AS TRABAJADORES/AS",
        "intro_suxeito_b": "{tipo_zona_cap} de {area} en el {centro}",
        "intro_texto_a": (
            "En base a la documentación y legislación de referencia se han determinado en el "
            "{centro} las siguientes zonas de trabajo susceptibles de exposición a radón por "
            "su situación en planta baja o bajo rasante, con la ocupación que se señala a "
            "continuación:"
        ),
        "intro_texto_a_sen_ocupacion": (
            "En base a la documentación y legislación de referencia se han determinado en el "
            "{centro} las siguientes zonas de trabajo susceptibles de exposición a radón por "
            "su situación en planta baja o bajo rasante:"
        ),
        "intro_texto": (
            "El {suxeito}, consta de los siguientes puestos de trabajo y la siguiente "
            "ocupación por espacio de trabajo, susceptibles de exposición a radón por su "
            "situación en planta baja o bajo rasante:"
        ),
        "sin_dato": ".....",
        "num_traballadores_texto": (
            "El número de trabajadores adscritos a este centro es de {total}, divididos en "
            "las siguientes categorías:"
        ),
        "cierre_grupo_a": "Los/as trabajadores/as del {centro} fueron informados de la ",
        "cierre_grupo_b": (
            "Los/as trabajadores/as {articulo} {tipo_zona} de {area} en el {centro}, fueron "
            "informados de la "
        ),
        "cierre_resto": (
            "realización de las mediciones y de su finalidad mediante {medio} el día {data}."
        ),
        "medio_defecto": "correo electrónico",
        "heading_4": "CONDICIONES DE LA EXPOSICIÓN",
        "texto_4": (
            "En el formulario de toma de datos del Anexo I se recogen: el tiempo de "
            "exposición de cada detector, el responsable de su colocación y retirada, la "
            "comprobación del estado de los detectores, las condiciones de temperatura o "
            "humedad, y observaciones sobre el sistema de ventilación."
        ),
        "texto_4_detectores": (
            "Descripción de los detectores utilizados: DETECTORES DE TRAZAS NUCLEARES EN "
            "ESTADO SOLIDO (SSNTD) DE LONGA DURACION. TIPO/MODELO: PADC/CR-39 – RSKS. "
            "FABRICANTE: RADOSYS-HUNGRIA."
        ),
        "heading_5": "PLANOS",
        "texto_5": (
            "El esquema gráfico del edificio y los planos de cada planta se muestran en el "
            "Anexo II, donde se indican las zonas de muestreo y las localizaciones de los "
            "detectores; junto con el código de cada zona de muestreo, se marca el código del "
            "detector correspondiente."
        ),
        "heading_7": "RESULTADOS DE LAS MEDICIONES REALIZADAS",
        "texto_7_intro": "Los resultados de las mediciones de radón realizadas pueden verse en la siguiente tabla:",
        "texto_7_sin_datos": "No hay datos de mediciones cargados para este centro de trabajo.",
        "tabla_headers": [
            "Código zona de muestreo",
            "Código detector",
            "Fecha de inicio exposición",
            "Fecha fin de exposición",
            "Concentración radón (Bq/m³) (*)",
            "Incertidumbre expandida y K",
            "Sala / Puesto de trabajo asociado",
        ],
        "header_sala_sin_puestos": "Sala del puesto de trabajo asociado",
        "tabla_nota": "* Se resaltan en negrita y letra roja los valores superiores a {nivel} Bq/m³.",
        "heading_8": "CONCLUSIONES",
        "conclusion_no_supera": (
            "A la vista de los datos del apartado 7, no hay ningún puesto de trabajo en el "
            "que se supere el nivel de referencia."
        ),
        "conclusion_supera": (
            "Las mediciones de radón en los siguientes puestos de trabajo: {salas}, dan "
            "valores superiores al nivel de referencia. Para estos casos se estudiarán "
            "medidas de remediación, se completarán las mediciones "
            "y/o se efectuarán mediciones en continuo, según proceda."
        ),
        "salas_indicadas": "indicados en el apartado 7",
        "heading_9": "DOCUMENTACIÓN DE REFERENCIA",
        "referencias": [
            "Ley 31/1995, de 8 de noviembre, de Prevención de Riesgos Laborales (BOE nº 269, de 10 de "
            "noviembre).",
            "Real Decreto 39/1997, de 17 de enero, por el que se aprueba el Reglamento de los Servicios "
            "de Prevención (BOE nº 27, de 31 de enero).",
            "Real Decreto 732/2019, de 20 de diciembre, por el que se modifica el Código Técnico de la "
            "Edificación, aprobado por el Real Decreto 314/2006, de 17 de marzo (BOE nº 311, de 27 de "
            "diciembre, páginas 140488 a 140674).",
            "Instrucción IS-47, de 9 de abril de 2025, del Consejo de Seguridad Nuclear por la que se "
            "aprueba el listado de términos municipales de actuación prioritaria frente al radón.",
            "Guía de Seguridad del CSN 11.4 Metodología para la evaluación de la exposición al radón en "
            "los lugares de trabajo.",
            "Estrategia para reducir la exposición al radón en Galicia. Estrategia Reduce Radón "
            "2025-2030. Consellería de Sanidade, Dirección Xeral de Saúde Pública, año 2024. "
            "https://airesaude.sergas.gal/Contidos/Documents/135/Estratexia%20reduce%20radon%202025-2030.pdf",
        ],
        "heading_10": "FECHA Y FIRMA DEL/DE LA TÉCNICO/A",
        "data_label_10": "Fecha",
        "firma_label": "Fdo.",
        "anexos": [
            "ANEXO I: FORMULARIOS DE TOMA DE DATOS",
            "ANEXO II: ESQUEMA GRÁFICO DEL EDIFICIO Y PLANOS DE CADA PLANTA",
            "ANEXO III: INFORME DE ENSAYO DEL LABORATORIO ACREDITADO",
            "ANEXO IV: CERTIFICADO ENAC DEL LABORATORIO ACREDITADO",
        ],
        "placeholder_centro": "...................",
        "placeholder_area": "...................",
    },
}


def _t(idioma: str, key: str):
    return TEXTOS.get(idioma, TEXTOS["gl"])[key]


@dataclass
class ReportContext:
    xerencia: str = ""
    cif: str = ""
    centro: str = ""
    enderezo: str = ""
    superficie_construida: str = ""
    superficie_util: str = ""
    num_plantas: str = ""
    postos_bullets: list = field(default_factory=list)
    num_traballadores_total: str = ""
    categorias_bullets: list = field(default_factory=list)
    notas_adicionais: str = ""
    texto_punto3_manual: str = ""  # si se rellena, sobrescribe todo el punto 3
    data_informacion_traballadores: str = ""
    medio_informacion_traballadores: str = "correo electrónico"
    data_informe: str = ""
    tecnico_nome: str = ""
    servizo_unidade: str = ""
    incertezas_por_defecto: str = ""
    conclusion_manual: str = ""
    # "unidade" | "servizo" | "outro" | "" (só se usa cando a área non é
    # ningunha das excepcións "xerais": determina como se nomea a área
    # nas frases que antes dicían sempre "unidade/servizo").
    tipo_zona: str = ""
    logo: bytes | None = None
    objeto_manual: str = ""  # si se rellena, sobrescribe el punto 2 (Objeto del informe)
    texto4_manual: str = ""  # si se rellena, sobrescribe el punto 4 (Condicións da exposición)
    texto5_manual: str = ""  # si se rellena, sobrescribe el punto 5 (Planos)


def _set_cell_shading(cell, color_hex: str):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    cell._tc.get_or_add_tcPr().append(shd)


def _set_cell_borders(cell, sz=4, color="000000"):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(sz))
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        borders.append(el)
    tcPr.append(borders)


def _set_cell_vertical_margins(cell, top_twips: int = 100, bottom_twips: int = 100):
    tcPr = cell._tc.get_or_add_tcPr()
    mar = tcPr.find(qn("w:tcMar"))
    if mar is None:
        mar = OxmlElement("w:tcMar")
        tcPr.append(mar)
    for edge, value in (("top", top_twips), ("bottom", bottom_twips)):
        el = mar.find(qn(f"w:{edge}"))
        if el is None:
            el = OxmlElement(f"w:{edge}")
            mar.append(el)
        el.set(qn("w:w"), str(value))
        el.set(qn("w:type"), "dxa")


def _set_table_fixed_layout(table, total_width_cm: float):
    tbl = table._tbl
    tblPr = tbl.tblPr
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tblPr.append(layout)
    dxa = int(total_width_cm * 566.929)
    tblW = tblPr.find(qn("w:tblW"))
    if tblW is None:
        tblW = OxmlElement("w:tblW")
        tblPr.append(tblW)
    tblW.set(qn("w:type"), "dxa")
    tblW.set(qn("w:w"), str(dxa))


def _fit_image_dims(image_bytes: bytes, max_width_cm: float, max_height_cm: float):
    try:
        img = Image.open(io.BytesIO(image_bytes))
        w_px, h_px = img.size
    except Exception:
        return Cm(max_width_cm), Cm(max_height_cm)

    aspect = w_px / h_px if h_px else 1
    max_width = Cm(max_width_cm)
    max_height = Cm(max_height_cm)

    height = max_height
    width = int(height * aspect)
    if width > max_width:
        width = max_width
        height = int(width / aspect)
    return width, height


def _add_field(paragraph, field_code: str, cached_text: str = "1", font_size_pt: float = 9):
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), field_code)

    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(int(font_size_pt * 2)))
    rpr.append(sz)
    run.append(rpr)
    t = OxmlElement("w:t")
    t.text = cached_text
    run.append(t)
    fld.append(run)

    paragraph._p.append(fld)


def _heading(doc, number: str, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(f"{number}. {text}")
    run.bold = True
    run.font.size = Pt(10)
    return p


def _body(doc, text: str):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    for run in p.runs:
        run.font.size = Pt(9.5)
    return p


def _set_font(doc):
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(9.5)


def _build_header(section, ctx: ReportContext, idioma: str):
    header = section.header
    header.is_linked_to_previous = False
    for p in list(header.paragraphs):
        p.text = ""

    table = header.add_table(rows=1, cols=2, width=Cm(17))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_fixed_layout(table, 17)
    LOGO_CELL_WIDTH_CM = 13.2
    DATE_CELL_WIDTH_CM = 3.8
    widths = [Cm(LOGO_CELL_WIDTH_CM), Cm(DATE_CELL_WIDTH_CM)]
    for col, w in zip(table.columns, widths):
        col.width = w
    for cell, w in zip(table.rows[0].cells, widths):
        cell.width = w
        _set_cell_borders(cell)

    logo_cell = table.rows[0].cells[0]
    logo_cell.vertical_alignment = 1
    para = logo_cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    logo_a_usar = ctx.logo or logo_por_defecto()
    if logo_a_usar:
        try:
            run = para.add_run()
            width, height = _fit_image_dims(
                logo_a_usar, max_width_cm=LOGO_CELL_WIDTH_CM - 0.6, max_height_cm=1.8
            )
            run.add_picture(io.BytesIO(logo_a_usar), width=width, height=height)
        except Exception:
            para.add_run("[logotipo]").italic = True
    else:
        r = para.add_run(_t(idioma, "logo_placeholder"))
        r.font.size = Pt(8)
        r.italic = True

    data_cell = table.rows[0].cells[1]
    data_cell.vertical_alignment = 1
    dp = data_cell.paragraphs[0]
    dp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r1 = dp.add_run(f"{_t(idioma, 'data_label')}: {ctx.data_informe}\n")
    r1.font.size = Pt(9)
    dp2 = data_cell.add_paragraph()
    r2 = dp2.add_run(f"{_t(idioma, 'paxina_label')} ")
    r2.font.size = Pt(9)
    _add_field(dp2, "PAGE", cached_text="1", font_size_pt=9)
    r3 = dp2.add_run(f" {_t(idioma, 'de_label')} ")
    r3.font.size = Pt(9)
    _add_field(dp2, "NUMPAGES", cached_text="1", font_size_pt=9)

    title_table = header.add_table(rows=1, cols=1, width=Cm(17))
    _set_table_fixed_layout(title_table, 17)
    cell = title_table.rows[0].cells[0]
    cell.vertical_alignment = 1
    _set_cell_borders(cell, sz=8)
    _set_cell_shading(cell, "F2F2F2")
    tp = cell.paragraphs[0]
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    centro_txt = ctx.centro.upper() if ctx.centro else _t(idioma, "placeholder_centro")
    if _area_grupo(ctx.servizo_unidade) == "A":
        titulo_texto = _t(idioma, "titulo_grupo_a").format(centro=centro_txt)
    else:
        area_txt = ctx.servizo_unidade.upper() if ctx.servizo_unidade else _t(idioma, "placeholder_area")
        tipo_zona_upper = _tipo_zona_texto(ctx.tipo_zona, idioma).upper()
        titulo_texto = _t(idioma, "titulo_grupo_b").format(area=area_txt, centro=centro_txt, tipo_zona_upper=tipo_zona_upper)
    tr = tp.add_run(titulo_texto)
    tr.bold = True
    tr.font.size = Pt(10.5)

    spacer = header.add_paragraph()
    spacer.paragraph_format.space_after = Pt(0)
    spacer_run = spacer.add_run("")
    spacer_run.font.size = Pt(6)


def _section_1(doc, ctx: ReportContext, idioma: str):
    _heading(doc, "1", _t(idioma, "heading_1"))
    table = doc.add_table(rows=5, cols=4)
    table.style = "Table Grid"

    cells0 = table.rows[0].cells
    cells0[0].text = _t(idioma, "lbl_xerencia")
    cells0[0].paragraphs[0].runs[0].bold = True
    merged01 = cells0[1].merge(cells0[2])
    merged01.text = ctx.xerencia
    cells0[3].text = f"{_t(idioma, 'lbl_cif')}: {ctx.cif}"

    cells1 = table.rows[1].cells
    cells1[0].text = _t(idioma, "lbl_centro")
    cells1[0].paragraphs[0].runs[0].bold = True
    merged1 = cells1[1].merge(cells1[2]).merge(cells1[3])
    merged1.text = ctx.centro

    cells2 = table.rows[2].cells
    cells2[0].text = _t(idioma, "lbl_servizo")
    cells2[0].paragraphs[0].runs[0].bold = True
    merged2b = cells2[1].merge(cells2[2]).merge(cells2[3])
    merged2b.text = ctx.servizo_unidade

    cells3 = table.rows[3].cells
    cells3[0].text = _t(idioma, "lbl_enderezo")
    cells3[0].paragraphs[0].runs[0].bold = True
    merged2 = cells3[1].merge(cells3[2]).merge(cells3[3])
    merged2.text = ctx.enderezo

    cells4 = table.rows[4].cells
    cells4[0].text = _t(idioma, "lbl_descricion")
    cells4[0].paragraphs[0].runs[0].bold = True
    cells4[1].text = f"{_t(idioma, 'lbl_superficie_construida')}: {ctx.superficie_construida}"
    cells4[2].text = f"{_t(idioma, 'lbl_superficie_util')}: {ctx.superficie_util}"
    cells4[3].text = f"{_t(idioma, 'lbl_num_plantas')}: {ctx.num_plantas}"

    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)

    for row_cells in (cells0, cells1, cells2, cells3, cells4):
        _set_cell_shading(row_cells[0], "F2F2F2")

    for row in table.rows:
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        row.height = Cm(0.9)
        for cell in row.cells:
            _set_cell_vertical_margins(cell, top_twips=120, bottom_twips=120)

    doc.add_paragraph()


def generar_texto_objeto_automatico(ctx: ReportContext, idioma: str = "gl") -> str:
    """Texto completo del punto 2 (Objeto do informe) tal como quedaría,
    para mostrarlo como valor por defecto editable en la app."""
    centro_txt = ctx.centro or _t(idioma, "sin_dato")
    if _area_grupo(ctx.servizo_unidade) == "A":
        lugar = centro_txt
    else:
        area_txt = ctx.servizo_unidade or _t(idioma, "sin_dato")
        tipo_zona_cap = _tipo_zona_texto(ctx.tipo_zona, idioma).capitalize()
        lugar = _t(idioma, "obxecto_lugar_b").format(area=area_txt, centro=centro_txt, tipo_zona_cap=tipo_zona_cap)
    return _t(idioma, "obxecto_texto").format(lugar=lugar, nivel=NIVEL_REFERENCIA_BQ_M3)


def _section_2(doc, ctx: ReportContext, idioma: str):
    _heading(doc, "2", _t(idioma, "heading_2"))
    _body(doc, ctx.objeto_manual.strip() or generar_texto_objeto_automatico(ctx, idioma))


def _bullet_list(doc, items: list[str]):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(item)
        r.font.size = Pt(9.5)


def _construir_bloques_punto3(ctx: ReportContext, idioma: str) -> list[tuple[str, object]]:
    """Devuelve el punto 3 como una lista ordenada de bloques
    [("body", texto) | ("bullets", [items])], usada tanto para escribirlo en
    el .docx como para mostrarlo como texto plano editable en la app."""
    centro_txt = ctx.centro or _t(idioma, "sin_dato")
    grupo = _area_grupo(ctx.servizo_unidade)
    if grupo == "A":
        suxeito = centro_txt
    else:
        area_txt = ctx.servizo_unidade or _t(idioma, "sin_dato")
        tipo_zona_cap = _tipo_zona_texto(ctx.tipo_zona, idioma).capitalize()
        suxeito = _t(idioma, "intro_suxeito_b").format(area=area_txt, centro=centro_txt, tipo_zona_cap=tipo_zona_cap)

    if grupo == "A":
        # Si no hay ningún profesional metido en ninguna sala (en la
        # ventana de detectores), no tiene sentido prometer "coa
        # ocupación que se sinala a continuación" y luego no enseñar
        # ninguna ocupación real; se usa la variante sin esa frase.
        hay_ocupacion = any(":" in b for b in (ctx.postos_bullets or []))
        clave_intro_a = "intro_texto_a" if hay_ocupacion else "intro_texto_a_sen_ocupacion"
        texto_intro = _t(idioma, clave_intro_a).format(centro=centro_txt)
    else:
        texto_intro = _t(idioma, "intro_texto").format(suxeito=suxeito)

    bloques: list[tuple[str, object]] = [
        ("body", texto_intro)
    ]
    bloques.append(("bullets", ctx.postos_bullets) if ctx.postos_bullets else ("body", _t(idioma, "sin_dato")))

    bloques.append(
        (
            "body",
            _t(idioma, "num_traballadores_texto").format(
                total=ctx.num_traballadores_total or _t(idioma, "sin_dato")
            ),
        )
    )
    bloques.append(
        ("bullets", ctx.categorias_bullets) if ctx.categorias_bullets else ("body", _t(idioma, "sin_dato"))
    )

    if ctx.notas_adicionais.strip():
        bloques.append(("body", ctx.notas_adicionais.strip()))

    if grupo == "A":
        cierre_texto = _t(idioma, "cierre_grupo_a").format(centro=centro_txt)
    else:
        area_txt = ctx.servizo_unidade or _t(idioma, "sin_dato")
        tipo_zona_txt = _tipo_zona_texto(ctx.tipo_zona, idioma)
        articulo_txt = _articulo_zona(ctx.tipo_zona, idioma)
        cierre_texto = _t(idioma, "cierre_grupo_b").format(
            area=area_txt, centro=centro_txt, tipo_zona=tipo_zona_txt, articulo=articulo_txt
        )
    bloques.append(
        (
            "body",
            cierre_texto
            + _t(idioma, "cierre_resto").format(
                medio=ctx.medio_informacion_traballadores or _t(idioma, "medio_defecto"),
                data=ctx.data_informacion_traballadores or _t(idioma, "sin_dato"),
            ),
        )
    )
    return bloques


def generar_texto_punto3_automatico(ctx: ReportContext, idioma: str = "gl") -> str:
    """Texto completo del punto 3 tal como quedaría en el informe, en texto
    plano (viñetas como líneas que empiezan por '- '), para mostrarlo como
    valor por defecto editable en la app."""
    lineas: list[str] = []
    for tipo, contenido in _construir_bloques_punto3(ctx, idioma):
        if tipo == "body":
            lineas.append(str(contenido))
        else:
            lineas.extend(f"- {item}" for item in contenido)
        lineas.append("")
    return "\n".join(lineas).strip()


def _section_3(doc, ctx: ReportContext, idioma: str):
    _heading(doc, "3", _t(idioma, "heading_3"))

    if ctx.texto_punto3_manual.strip():
        for linea in ctx.texto_punto3_manual.strip().splitlines():
            linea = linea.strip()
            if not linea:
                continue
            if linea.startswith(("-", "•")):
                p = doc.add_paragraph(style="List Bullet")
                r = p.add_run(linea.lstrip("-•").strip())
                r.font.size = Pt(9.5)
            else:
                _body(doc, linea)
        return

    for tipo, contenido in _construir_bloques_punto3(ctx, idioma):
        if tipo == "body":
            _body(doc, contenido)
        else:
            _bullet_list(doc, contenido)


def generar_texto4_automatico(idioma: str = "gl") -> str:
    """Texto completo del punto 4 (Condicións da exposición), para
    mostrarlo como valor por defecto editable en la app."""
    return (_t(idioma, "texto_4") + "\n\n" + _t(idioma, "texto_4_detectores")).strip()


def _section_4(doc, ctx: ReportContext, idioma: str):
    _heading(doc, "4", _t(idioma, "heading_4"))
    if ctx.texto4_manual.strip():
        _body(doc, ctx.texto4_manual)
        return
    _body(doc, _t(idioma, "texto_4"))
    _body(doc, _t(idioma, "texto_4_detectores"))


def generar_texto5_automatico(idioma: str = "gl") -> str:
    """Texto completo del punto 5 (Planos), para mostrarlo como valor
    por defecto editable en la app."""
    return _t(idioma, "texto_5")


def _section_5(doc, ctx: ReportContext, idioma: str):
    _heading(doc, "5", _t(idioma, "heading_5"))
    if ctx.texto5_manual.strip():
        _body(doc, ctx.texto5_manual)
        return
    _body(doc, _t(idioma, "texto_5"))


def _row_value(row, *keys, default=""):
    for k in keys:
        if k in row and pd.notna(row[k]):
            return row[k]
    return default


def _quitar_turno_de_puestos(texto: str) -> str:
    """En la tabla del punto 7 no deben figurar los horarios de los
    trabajadores: quita el " - Turno" de cada categoría de
    "Profesionales en la sala" (p.ej. "PSG (1) - Mañana, DUE (1) -
    Tarde" -> "PSG (1), DUE (1)"), dejando solo la categoría y el
    número de personas. Si el texto no tiene turno (formato antiguo),
    se deja tal cual."""
    if not texto:
        return texto
    partes = [p.strip() for p in re.split(r"[,\n]", texto) if p.strip()]
    partes_sin_turno = [p.split(" - ", 1)[0].strip() for p in partes]
    return ", ".join(partes_sin_turno)


def _section_7(doc, df: pd.DataFrame, ctx: ReportContext, idioma: str):
    _heading(doc, "7", _t(idioma, "heading_7"))
    _body(doc, _t(idioma, "texto_7_intro"))

    headers = list(_t(idioma, "tabla_headers"))
    tiene_puestos_datos = False
    for _, row in df.iterrows():
        puestos_chk = _row_value(row, "Profesionales en la sala", "Puestos en la sala", default="")
        if pd.notna(puestos_chk) and str(puestos_chk).strip():
            tiene_puestos_datos = True
            break
    if not tiene_puestos_datos:
        headers[-1] = _t(idioma, "header_sala_sin_puestos")

    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        _set_cell_shading(hdr_cells[i], "D9D9D9")
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            r.bold = True
            r.font.size = Pt(8)

    any_exceeds = False
    for _, row in df.iterrows():
        codigo_zona = _row_value(row, "Código de la sala", "Código Sala", "Sala")
        codigo_zona = "" if pd.isna(codigo_zona) else str(codigo_zona).rstrip(".0")
        codigo_det = _row_value(row, "Código")
        fecha_ini = _row_value(row, "Fecha de colocación fmt")
        fecha_fin = _row_value(row, "Fecha de retirada real fmt")
        concentracion = row.get("Resultado Bq/m3")
        sala_nome = _row_value(row, "Sala", default="")
        sala_txt = str(sala_nome).strip() if pd.notna(sala_nome) and str(sala_nome).strip() else ""
        puestos = _row_value(row, "Profesionales en la sala", "Puestos en la sala", default="")
        if pd.notna(puestos) and not isinstance(puestos, str):
            puestos = str(int(puestos)) if float(puestos).is_integer() else str(puestos)
        puestos_txt = _quitar_turno_de_puestos(str(puestos).strip()) if puestos else ""

        incerteza = row.get("Incerteza expandida e K", ctx.incertezas_por_defecto)
        if pd.isna(incerteza) or incerteza is None:
            incerteza = ctx.incertezas_por_defecto or ""
        incerteza_txt = str(incerteza).strip()
        if incerteza_txt and not incerteza_txt.startswith(("±", "+/-", "+-")):
            incerteza_txt = f"±{incerteza_txt}"

        cells = table.add_row().cells
        values = [
            str(codigo_zona) if codigo_zona else "",
            str(codigo_det) if pd.notna(codigo_det) else "",
            str(fecha_ini),
            str(fecha_fin),
            f"{concentracion:g}" if pd.notna(concentracion) else "",
            incerteza_txt,
        ]
        exceeds = pd.notna(concentracion) and concentracion > NIVEL_REFERENCIA_BQ_M3
        any_exceeds = any_exceeds or exceeds

        for i, val in enumerate(values):
            cells[i].text = val
            cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.font.size = Pt(8.5)
                if exceeds and i == 4:
                    r.bold = True
                    r.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

        last_cell = cells[6]
        last_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        last_cell.text = ""
        p1 = last_cell.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = p1.add_run(sala_txt)
        r1.font.size = Pt(8.5)
        r1.bold = True
        if puestos_txt:
            p2 = last_cell.add_paragraph()
            p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r2 = p2.add_run(puestos_txt)
            r2.font.size = Pt(8.5)

    note = doc.add_paragraph()
    note.paragraph_format.space_before = Pt(4)
    nr = note.add_run(_t(idioma, "tabla_nota").format(nivel=NIVEL_REFERENCIA_BQ_M3))
    nr.italic = True
    nr.font.size = Pt(8)
    return any_exceeds


def _unir_con_conjuncion(items: list[str], idioma: str) -> str:
    """Une una lista de elementos con comas, y con la conjunción antes
    del último cuando hay más de uno ("y" en castellán, "e" en
    galego): "Sala 1", "Sala 1 e Sala 2", "Sala 1, Sala 2 e Sala 3"."""
    items = [i for i in items if i]
    if not items:
        return ""
    if len(items) == 1:
        return items[0]
    conjuncion = "e" if idioma == "gl" else "y"
    return ", ".join(items[:-1]) + f" {conjuncion} {items[-1]}"


def generar_conclusion_automatica(df: pd.DataFrame, idioma: str = "gl") -> str:
    if df is None or df.empty or "Resultado Bq/m3" not in df.columns:
        return _t(idioma, "conclusion_no_supera")
    mask = df["Resultado Bq/m3"] > NIVEL_REFERENCIA_BQ_M3
    if not mask.any():
        return _t(idioma, "conclusion_no_supera")
    exceeded_rooms: list[str] = []
    for _, row in df[mask].iterrows():
        sala = _row_value(row, "Sala", "Código de la sala", "Código Sala", default="")
        if sala:
            exceeded_rooms.append(str(sala))
    salas = _unir_con_conjuncion(exceeded_rooms, idioma) if exceeded_rooms else _t(idioma, "salas_indicadas")
    return _t(idioma, "conclusion_supera").format(salas=salas)


def _section_8(doc, ctx: ReportContext, any_exceeds: bool, exceeded_rooms: list[str], idioma: str):
    _heading(doc, "8", _t(idioma, "heading_8"))
    if ctx.conclusion_manual.strip():
        _body(doc, ctx.conclusion_manual.strip())
        return
    if not any_exceeds:
        _body(doc, _t(idioma, "conclusion_no_supera"))
    else:
        salas = _unir_con_conjuncion(exceeded_rooms, idioma) if exceeded_rooms else _t(idioma, "salas_indicadas")
        _body(doc, _t(idioma, "conclusion_supera").format(salas=salas))


def _section_9(doc, idioma: str):
    _heading(doc, "9", _t(idioma, "heading_9"))
    for ref in _t(idioma, "referencias"):
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(ref)
        run.font.size = Pt(9)


def _section_10(doc, ctx: ReportContext, idioma: str):
    p_heading = _heading(doc, "10", _t(idioma, "heading_10"))
    p_fecha = _body(doc, f"{_t(idioma, 'data_label_10')}: {ctx.data_informe or _t(idioma, 'sin_dato')}")
    p_blank1 = doc.add_paragraph()
    p_blank2 = doc.add_paragraph()
    _body(doc, f"{_t(idioma, 'firma_label')}: {ctx.tecnico_nome or _t(idioma, 'sin_dato')}")
    # Que este bloque no se corte entre dúas páxinas: non se permite
    # salto de páxina entre ningún destes parágrafos e o seguinte.
    for p in (p_heading, p_fecha, p_blank1, p_blank2):
        p.paragraph_format.keep_with_next = True


def _anexos(doc, idioma: str):
    doc.add_paragraph()
    for a in _t(idioma, "anexos"):
        p = doc.add_paragraph()
        r = p.add_run(a)
        r.bold = True
        r.font.size = Pt(9.5)


def _force_field_update_on_open(doc):
    settings = doc.settings.element
    update_fields = OxmlElement("w:updateFields")
    update_fields.set(qn("w:val"), "true")
    settings.append(update_fields)


def generate_report(ctx: ReportContext, df: pd.DataFrame, idioma: str = "gl") -> io.BytesIO:
    """Genera el informe .docx y lo devuelve como BytesIO listo para descargar.

    `idioma`: "gl" (gallego, por defecto, idioma oficial de la plantilla) o
    "es" (castellano). El contenido dinamico de `ctx` (postos_bullets,
    categorias_bullets) debe venir ya en el idioma deseado."""
    if idioma not in TEXTOS:
        idioma = "gl"

    doc = Document()
    section = doc.sections[0]
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)
    section.top_margin = Cm(4.0)
    section.bottom_margin = Cm(2)
    section.header_distance = Cm(0.8)
    _set_font(doc)
    _build_header(section, ctx, idioma)
    _force_field_update_on_open(doc)

    _section_1(doc, ctx, idioma)
    _section_2(doc, ctx, idioma)
    _section_3(doc, ctx, idioma)
    _section_4(doc, ctx, idioma)
    _section_5(doc, ctx, idioma)
    any_exceeds = False
    exceeded_rooms: list[str] = []
    if not df.empty:
        any_exceeds = _section_7(doc, df, ctx, idioma)
        if any_exceeds:
            mask = df["Resultado Bq/m3"] > NIVEL_REFERENCIA_BQ_M3
            for _, row in df[mask].iterrows():
                sala = _row_value(row, "Sala", "Código de la sala", "Código Sala", default="")
                if sala:
                    exceeded_rooms.append(str(sala))
    else:
        _heading(doc, "7", _t(idioma, "heading_7"))
        _body(doc, _t(idioma, "texto_7_sin_datos"))

    _section_8(doc, ctx, any_exceeds, exceeded_rooms, idioma)
    _section_9(doc, idioma)
    _section_10(doc, ctx, idioma)
    _anexos(doc, idioma)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
